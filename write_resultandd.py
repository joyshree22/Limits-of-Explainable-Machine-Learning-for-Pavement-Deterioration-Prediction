"""
Generate resultandd.docx from the corrected primary-analysis outputs.

The document intentionally follows the article/methodology spine:
benchmarks, design vs monitoring, persistence comparison, LOO transfer,
and cautious interpretability. It does not report post-hoc regime/routing
experiments as primary findings.
"""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).parent
RES = ROOT / "results"
FIG = ROOT / "figures"
OUT = ROOT / "resultandd.docx"


def read_csv(name: str) -> pd.DataFrame:
    path = RES / name
    if not path.exists():
        raise FileNotFoundError(f"Missing required result file: {path}")
    return pd.read_csv(path)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, title: str, df: pd.DataFrame, columns: list[str]) -> None:
    add_para(doc, title)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr[i].text = col
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            val = row.get(col, "")
            if isinstance(val, float):
                cells[i].text = f"{val:.4g}"
            else:
                cells[i].text = str(val)
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str, width: float = 6.2) -> None:
    path = FIG / filename
    if not path.exists():
        add_para(doc, f"[Figure missing: {filename}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(9)


def metric(df: pd.DataFrame, model: str, col: str) -> float:
    rows = df[df["model"] == model]
    return float(rows.iloc[0][col]) if len(rows) else float("nan")


def main() -> None:
    test = read_csv("test_metrics.csv")
    bench = read_csv("benchmark_metrics.csv")
    common = read_csv("common_set_comparison.csv")
    loo = read_csv("loo_summary.csv")
    shap = read_csv("shap_consistency.csv")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    add_heading(doc, "4. Results and Discussion", 1)
    add_para(
        doc,
        "This section reports the corrected primary analysis after removing same-visit "
        "condition-measure leakage, excluding geographic/region proxy features, fixing "
        "section-level bootstrap resampling, and recomputing monitoring-vs-persistence "
        "comparisons on identical section/date rows. The results should be read as a "
        "stress test of transferability, not as an exercise in maximizing R2."
    )

    add_heading(doc, "4.1 Benchmark Performance", 2)
    add_para(
        doc,
        "Mean, AASHTO age-only, Ridge, and persistence benchmarks establish the "
        "reference frame for the ensemble models. Ridge no longer produces near-perfect "
        "distress performance after the distress sub-indicator leakage was removed."
    )
    add_table(
        doc,
        "Table 1. Benchmark model performance.",
        bench,
        ["model", "n", "R2", "RMSE", "MAE"],
    )

    add_heading(doc, "4.2 Ensemble Test-Set Performance", 2)
    add_para(
        doc,
        "Design-task performance remains modest and uncertain across targets. Monitoring "
        "improves pooled test metrics, but this improvement must be interpreted against "
        "persistence because lagged condition is highly informative."
    )
    add_table(
        doc,
        "Table 2. Corrected test-set performance with section-bootstrap R2 intervals.",
        test,
        ["model", "n_obs", "n_sections", "R2", "R2_CI_lower", "R2_CI_upper", "RMSE", "MAE"],
    )

    add_heading(doc, "4.3 Monitoring Versus Persistence", 2)
    add_para(
        doc,
        "The common-set table compares trained monitoring models and persistence on the "
        "same held-out section/date rows. Absolute IRI monitoring underperforms one-year "
        "persistence on this common set, while the delta-IRI sensitivity model improves "
        "absolute-scale RMSE after adding predicted deterioration increments back to the lag."
    )
    add_table(
        doc,
        "Table 3. Common-set monitoring and persistence comparison.",
        common,
        ["model", "n", "R2", "RMSE", "MAE", "skill_vs_persist_k1"],
    )

    add_heading(doc, "4.4 Leave-One-Region-Out Generalization", 2)
    ont_xgb = loo[(loo["withheld_region"] == "Ontario") & (loo["arch"] == "xgb")]
    ont_r2 = float(ont_xgb.iloc[0]["R2"]) if len(ont_xgb) else float("nan")
    add_para(
        doc,
        f"LOO validation again shows weak climate transfer. Ontario XGBoost LOO R2 is "
        f"{ont_r2:.3f}, so the corrected primary run does not support the earlier "
        "post-hoc claim that Ontario transfer became positive. Georgia also remains a "
        "major failure case, showing that climate mismatch is not captured by freeze "
        "index alone."
    )
    add_table(
        doc,
        "Table 4. Leave-one-region-out IRI design validation.",
        loo,
        ["withheld_region", "arch", "target", "n_obs", "R2", "RMSE", "MAE"],
    )
    add_figure(
        doc,
        "loo_vs_climate_gradient.png",
        "Figure 1. LOO R2 across the climate gradient.",
    )

    add_heading(doc, "4.5 Interpretability Stability", 2)
    rho = float(shap.iloc[0]["spearman_rho"])
    threshold = float(shap.iloc[0]["threshold"])
    add_para(
        doc,
        f"Cross-model SHAP rank consistency is rho = {rho:.3f}, below the "
        f"pre-specified stability threshold of {threshold:.2f}. SHAP and PDP outputs "
        "therefore describe model behavior on this small dataset; robust claims should "
        "be restricted to converging top-ranked features and should not be treated as "
        "causal pavement-mechanism evidence."
    )
    add_table(
        doc,
        "Table 5. SHAP cross-model consistency.",
        shap,
        ["n_features", "spearman_rho", "p_value", "threshold", "status"],
    )
    add_figure(
        doc,
        "shap_bar_iri_design.png",
        "Figure 2. Global SHAP importance for the corrected IRI design model.",
        width=5.4,
    )
    add_figure(
        doc,
        "residual_diagnostics.png",
        "Figure 3. Residual diagnostics for the corrected XGBoost IRI design model.",
    )

    add_heading(doc, "4.6 Synthesis", 2)
    add_para(
        doc,
        "The corrected pipeline supports the article's central caution. High apparent "
        "accuracy can arise from leakage, lag persistence, regional imbalance, or proxy "
        "features. After removing those shortcuts, design-stage transfer remains weak, "
        "monitoring must be benchmarked against persistence, and interpretability is "
        "bounded by model generalization and cross-model stability."
    )

    doc.save(OUT)
    print(f"Saved -> {OUT}")


if __name__ == "__main__":
    main()
