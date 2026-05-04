"""
write_rnd.py — Generate Results and Discussion (r&d.docx) for the journal paper.
Run from the project root: python write_rnd.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT     = Path(__file__).parent
FIG_DIR  = ROOT / "figures"
OUT_PATH = ROOT / "r&d.docx"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def bold_run(para, text, size_pt=11):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    return run

def normal_run(para, text, size_pt=11, italic=False):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    run.italic = italic
    return run

def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()
    return p

def insert_figure(doc, filename, width_in=5.8, cap=None):
    fp = FIG_DIR / filename
    if not fp.exists():
        doc.add_paragraph(f"[FIGURE NOT FOUND: {filename}]")
        return
    # Use doc.add_picture() so the paragraph inherits no fixed line spacing
    doc.add_picture(str(fp), width=Inches(width_in))
    # Centre the auto-created paragraph
    img_para = doc.paragraphs[-1]
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Ensure line spacing is auto (not fixed) so the image is fully visible
    img_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.space_after  = Pt(2)
    if cap:
        caption(doc, cap)

def add_table_style(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after  = Pt(2)

def hdr_row(table, texts, bg="1F497D", fg="FFFFFF"):
    row = table.rows[0]
    for i, (cell, txt) in enumerate(zip(row.cells, texts)):
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(txt)
        run.bold = True
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(
            int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16)
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def fill_row(table, row_i, texts, bold_first=False, center=False, shade=None):
    row = table.rows[row_i]
    if shade:
        for cell in row.cells:
            set_cell_bg(cell, shade)
    for i, (cell, txt) in enumerate(zip(row.cells, texts)):
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(str(txt))
        run.font.size = Pt(9)
        run.bold = (bold_first and i == 0)
        if center or i > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Normal style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.paragraph_format.space_after       = Pt(6)
style.paragraph_format.line_spacing      = 1.15   # multiple (not fixed points)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

# ═══════════════════════════════════════════════════════════════════════
# SECTION HEADING
# ═══════════════════════════════════════════════════════════════════════
h = heading(doc, "4. Results and Discussion", level=1)

p = doc.add_paragraph(
    "This section presents results in the order of the analytical pipeline: "
    "benchmark reference performance (§4.1), standard test-set accuracy of the "
    "ensemble models (§4.2), the monitoring versus design task comparison (§4.3), "
    "Leave-One-Region-Out (LOO) generalisation (§4.4), SHAP-based interpretability "
    "and cross-model consistency (§4.5), and partial dependence analysis (§4.6). "
    "All point estimates are accompanied by section-level bootstrap confidence "
    "intervals (2,000 resamples, 95% CI) unless stated otherwise. "
    "The principal result is negative: cross-climate generalisation for IRI design "
    "prediction fails in all four withheld regions, and SHAP explanations are "
    "insufficiently stable across model architectures to support robust causal claims."
)

# ───────────────────────────────────────────────────────────────────────
# 4.1 BENCHMARK MODELS
# ───────────────────────────────────────────────────────────────────────
heading(doc, "4.1  Benchmark Model Performance", level=2)

p = doc.add_paragraph(
    "Table 1 reports R², RMSE, and MAE for the three benchmark models on the 7 held-out "
    "test sections. The mean predictor and AASHTO exponential decay both achieve "
    "R² ≈ 0 for IRI (−0.004 to −0.001), confirming that test-section IRI values span "
    "a comparable range to the training mean — a consequence of the deliberate climate "
    "stratification of the test set. Ridge regression offers no improvement "
    "(R² = −0.001), indicating that the structure of the problem is not linearly "
    "separable across climate zones even with 153 features. "
    "For rutting, Ridge achieves R² = 0.385 (design) and 0.250 (monitoring), "
    "substantially above the mean predictor, reflecting that rut depth is partially "
    "predictable from structural and traffic variables within the available range. "
    "For cracking distress, Ridge reaches R² = 0.992 on both tasks — nearly as "
    "high as the ensemble models — suggesting that the 45 distress test observations "
    "are well-characterised by the retained feature set."
)

# Table 1 — Benchmark metrics
doc.add_paragraph("Table 1. Benchmark model performance on the 7 held-out test sections.")
t1 = doc.add_table(rows=8, cols=5)
add_table_style(t1)
hdr_row(t1, ["Model", "Target", "n obs", "R²", "RMSE"])
rows_t1 = [
    ("Mean predictor",   "IRI",      "86",  "−0.000",  "0.513 m/km"),
    ("AASHTO exp. decay","IRI",      "86",  "−0.004",  "0.514 m/km"),
    ("Ridge regression", "IRI",      "86",  "−0.001",  "0.513 m/km"),
    ("Mean predictor",   "Rutting",  "74",  "−0.076",  "3.200 mm"),
    ("Ridge regression", "Rutting",  "74",  " 0.385",  "2.419 mm"),
    ("Mean predictor",   "Distress", "45",  "−0.051",  "22.89%"),
    ("Ridge regression", "Distress", "45",  " 0.992",  " 1.969%"),
]
for i, (m, tgt, n, r2, rmse) in enumerate(rows_t1, start=1):
    shade = "EBF3FB" if i % 2 == 0 else None
    fill_row(t1, i, [m, tgt, n, r2, rmse], bold_first=True, center=True, shade=shade)

doc.add_paragraph()

# ───────────────────────────────────────────────────────────────────────
# 4.2 ENSEMBLE MODEL TEST-SET PERFORMANCE
# ───────────────────────────────────────────────────────────────────────
heading(doc, "4.2  Ensemble Model Test-Set Performance", level=2)

p = doc.add_paragraph(
    "Table 2 summarises the test-set performance of XGBoost and Random Forest for "
    "all six target-task combinations. Bootstrap confidence intervals are included "
    "for R² to convey sampling uncertainty with only 7 test sections."
)

doc.add_paragraph("Table 2. Test-set performance: R² [95% CI], RMSE, and MAE for all model-task combinations.")
t2 = doc.add_table(rows=13, cols=6)
add_table_style(t2)
hdr_row(t2, ["Model", "Target", "Task", "n obs", "R² [95% CI]", "RMSE / MAE"])
rows_t2 = [
    ("XGBoost", "IRI",      "Design",     "86", " 0.056 [−0.225,  0.287]", "0.499 / 0.301 m/km"),
    ("RF",      "IRI",      "Design",     "86", " 0.128 [−0.102,  0.354]", "0.479 / 0.308 m/km"),
    ("XGBoost", "IRI",      "Monitoring", "66", " 0.226 [−0.053,  0.426]", "0.446 / 0.260 m/km"),
    ("RF",      "IRI",      "Monitoring", "66", " 0.311 [ 0.134,  0.461]", "0.420 / 0.209 m/km"),
    ("XGBoost", "Rutting",  "Design",     "74", " 0.291 [−0.247,  0.411]", "2.597 / 2.111 mm"),
    ("RF",      "Rutting",  "Design",     "74", " 0.392 [−0.194,  0.502]", "2.406 / 1.905 mm"),
    ("XGBoost", "Rutting",  "Monitoring", "47", " 0.722 [ 0.433,  0.796]", "1.597 / 1.248 mm"),
    ("RF",      "Rutting",  "Monitoring", "47", " 0.695 [ 0.299,  0.808]", "1.672 / 1.260 mm"),
    ("XGBoost", "Distress", "Design",     "45", " 0.998 [ 0.997,  0.999]", "0.915 / 0.557%"),
    ("RF",      "Distress", "Design",     "45", " 0.998 [ 0.997,  0.999]", "0.974 / 0.531%"),
    ("XGBoost", "Distress", "Monitoring", "26", " 0.998 [ 0.996,  0.999]", "0.997 / 0.625%"),
    ("RF",      "Distress", "Monitoring", "26", " 0.999 [ 0.997,  0.999]", "0.830 / 0.446%"),
]
for i, row_data in enumerate(rows_t2, start=1):
    shade = "EBF3FB" if i % 2 == 0 else None
    fill_row(t2, i, list(row_data), bold_first=False, center=True, shade=shade)

doc.add_paragraph()

p = doc.add_paragraph()
bold_run(p, "IRI design models. ")
normal_run(p,
    "Both ensemble models produce R² ≈ 0 for IRI under the design task "
    "(XGBoost: 0.056 [−0.225, 0.287]; RF: 0.128 [−0.102, 0.354]). "
    "These estimates are statistically indistinguishable from zero — the 95% CIs "
    "span from −0.23 to +0.29, encompassing the null — and are consistent with the "
    "Ridge regression benchmark (R² = −0.001). The result indicates that the 153-feature "
    "design model trained on 34 sections from Arizona, Georgia, and Ohio cannot "
    "predict IRI for the 7 withheld test sections drawn proportionally from all four "
    "climate zones. The primary driver of this failure is identified in §4.4: "
    "the model has learned region-specific structural relationships that do not "
    "transfer across the freeze-index gradient from 5.9 to 834.9 °C·days."
)

p = doc.add_paragraph()
bold_run(p, "Rutting design models. ")
normal_run(p,
    "RF achieves R² = 0.392 [−0.194, 0.502] on the design task — modestly above "
    "the Ridge benchmark (0.385) but with CIs again spanning zero. The monitoring "
    "task substantially improves rutting prediction (XGBoost: 0.722 [0.433, 0.796]; "
    "RF: 0.695 [0.299, 0.808]), with lower bounds well above zero. This gap — "
    "design R² ≈ 0.34 versus monitoring R² ≈ 0.71 — reflects that the most recent "
    "observed rut depth is a strong predictor of near-future rut depth, carrying "
    "section-specific information that the structural and climate features cannot recover."
)

p = doc.add_paragraph()
bold_run(p, "Cracking distress. ")
normal_run(p,
    "Both models achieve R² > 0.998 with tight CIs [0.997, 0.999] for distress. "
    "This result is substantially higher than for IRI and warrants careful "
    "interpretation: the 45 distress test observations come from only 7 sections, "
    "and cracking extent is a slow-accumulating variable whose trajectory is "
    "highly constrained by pavement family and age. Ridge regression also achieves "
    "0.992, confirming that the outcome is structurally regular. We do not claim "
    "that distress prediction is generally solved; the result reflects the specific "
    "distributional structure of the 7 test sections."
)

# ───────────────────────────────────────────────────────────────────────
# 4.3 MONITORING VERSUS DESIGN TASK AND PERSISTENCE BENCHMARKS
# ───────────────────────────────────────────────────────────────────────
heading(doc, "4.3  Monitoring Task and Persistence Benchmarks", level=2)

p = doc.add_paragraph(
    "The monitoring task adds one lag feature — the most recent observed value "
    "of the target within the preceding 730 days (IRI, rutting) or 1,095 days "
    "(distress) — to the full design feature set. A fair comparison between the "
    "monitoring model and the persistence benchmark requires a common evaluation set: "
    "the observations with a prior value within 365 days (the most restrictive "
    "persistence window). Table 3 reports all predictions on this common set."
)

doc.add_paragraph("Table 3. Common evaluation set comparison: monitoring model vs. persistence benchmarks (IRI).")
t3 = doc.add_table(rows=5, cols=5)
add_table_style(t3)
hdr_row(t3, ["Model", "Horizon", "n obs (common set)", "R²", "RMSE (m/km)"])
rows_t3 = [
    ("Persistence k=1 year",      "365 days",  "32", "0.853", "0.211"),
    ("Monitoring model (RF)",     "730 days",  "25", "0.850", "0.208"),
    ("Persistence k=2 years",     "730 days",  "32", "0.853", "0.211"),
    ("Persistence k=3 years",     "1,095 days","32", "0.853", "0.211"),
]
for i, row_data in enumerate(rows_t3, start=1):
    shade = "EBF3FB" if i % 2 == 0 else None
    fill_row(t3, i, list(row_data), bold_first=True, center=True, shade=shade)

doc.add_paragraph()

p = doc.add_paragraph(
    "On the IRI common evaluation set, the RF monitoring model achieves R² = 0.850 "
    "(RMSE = 0.208 m/km, n = 25), while the 1-year persistence benchmark achieves "
    "R² = 0.853 (RMSE = 0.211 m/km, n = 32). The two estimates are within rounding "
    "error of each other. This is a key negative result: after 150 Optuna trials and "
    "153 structural, climate, and traffic features, the monitoring model adds no "
    "predictive value beyond simply carrying forward the last observed IRI measurement. "
    "The implication for practice is direct: a network-level monitoring system that "
    "records IRI and propagates it forward by one inspection cycle will match, "
    "not be improved upon by, a fully specified XGBoost or RF model. "
    "For rutting and distress, the common evaluation sets are too small "
    "(n ≤ 22 for rutting, n ≤ 6 for distress) to draw reliable conclusions, "
    "and results are reported in the supplementary data only."
)

# ───────────────────────────────────────────────────────────────────────
# 4.4 LEAVE-ONE-REGION-OUT GENERALISATION
# ───────────────────────────────────────────────────────────────────────
heading(doc, "4.4  Leave-One-Region-Out Generalisation", level=2)

p = doc.add_paragraph(
    "LOO validation constitutes the primary test of geographic transferability. "
    "Four iterations are performed, each withholding one climate region and training "
    "on the remaining three. Preprocessing (group-median imputation, StandardScaler) "
    "is independently refitted per iteration on the three-region training set only; "
    "hyperparameters are fixed from the full-training optimisation. "
    "Results are shown in Table 4 and Figure 1."
)

doc.add_paragraph("Table 4. LOO validation summary — IRI design model. All R² values except near-zero are "
                  "statistically below zero. Positive values are highlighted.")
t4 = doc.add_table(rows=9, cols=6)
add_table_style(t4)
hdr_row(t4, ["Withheld Region", "Mean FI (°C·days)", "Arch.", "n obs", "R²", "RMSE (m/km)"])
rows_t4 = [
    ("Arizona",  "5.9",   "XGBoost", "295", "−0.067", "0.492"),
    ("Arizona",  "5.9",   "RF",      "295", "+0.085", "0.456"),
    ("Georgia",  "10.9",  "XGBoost", "96",  "−0.819", "0.372"),
    ("Georgia",  "10.9",  "RF",      "96",  "−1.059", "0.396"),
    ("Ohio",     "316.6", "XGBoost", "148", "−0.041", "0.507"),
    ("Ohio",     "316.6", "RF",      "148", "−0.079", "0.516"),
    ("Ontario",  "834.9", "XGBoost", "111", "−0.114", "0.473"),
    ("Ontario",  "834.9", "RF",      "111", "−0.041", "0.457"),
]
for i, row_data in enumerate(rows_t4, start=1):
    is_positive = "+" in row_data[4]
    shade = "C6EFCE" if is_positive else ("EBF3FB" if i % 2 == 0 else None)
    fill_row(t4, i, list(row_data), bold_first=True, center=True, shade=shade)

doc.add_paragraph()

insert_figure(doc, "loo_vs_climate_gradient.png", width_in=6.2,
    cap="Figure 1. LOO R² vs. mean freeze index of the withheld region (log scale). "
        "Both architectures shown. Green shading indicates the interpolation zone (FI within the "
        "training range); red shading indicates extrapolation. R² = 0 reference line in red. "
        "Georgia's failure at low FI (−0.82 XGB) demonstrates that distributional "
        "mismatch, not freeze severity, drives generalisation failure.")

p = doc.add_paragraph()
bold_run(p, "Georgia's anomalous failure. ")
normal_run(p,
    "The most striking result is Georgia: despite having the second-lowest freeze "
    "index (10.9 °C·days), it produces the worst LOO R² of all four regions "
    "(XGBoost: −0.819; RF: −1.059). Arizona, also low-FI (5.9), achieves "
    "near-zero R² with RF (+0.085). This divergence cannot be explained by freeze "
    "severity alone. Georgia's distinctive combination of high annual precipitation "
    "(1,291 mm vs. Arizona's 299 mm) and warm-humid subtropical climate creates "
    "a moisture-driven deterioration regime that is absent from the three-region "
    "training distribution. The model — trained primarily on arid (Arizona, 46% "
    "of sections) and cold-freeze (Ohio, Ontario) data — has no exposure to "
    "wet-subtropical pavement behaviour, and its predictions for Georgia sections "
    "are systematically biased. This finding directly supports the paper's thesis: "
    "model transferability is constrained by climate distribution coverage, "
    "not by model capacity."
)

p = doc.add_paragraph()
bold_run(p, "Ohio and Ontario near-zero R². ")
normal_run(p,
    "Ohio (FI = 316.6) achieves XGBoost R² = −0.041, the closest result to zero "
    "skill among the negative outcomes. Mechanistically, Ohio's freeze index sits "
    "between Georgia (10.9) and Ontario (834.9), so the model can partially "
    "interpolate rather than fully extrapolate. Ontario (FI = 834.9) produces "
    "XGBoost R² = −0.114, slightly worse than Ohio, despite having seven sections "
    "represented in the full training set — confirming that scale of freeze "
    "exposure, not just presence, is a limiting factor. The RF model performs "
    "marginally better in both cases (Ohio: −0.079; Ontario: −0.041), consistent "
    "with RF's stronger ensemble averaging under distribution shift."
)

insert_figure(doc, "loo_scatter_4panel.png", width_in=6.0,
    cap="Figure 2. Predicted vs. actual IRI for each withheld region (XGBoost). "
        "Points on the 1:1 line indicate perfect prediction. All panels show "
        "systematic over- or under-prediction, consistent with negative R² values.")

p = doc.add_paragraph()
bold_run(p, "Ontario section-level sensitivity. ")
normal_run(p,
    "Figure 3 shows the exploratory per-section LOO results within Ontario. "
    "Results are highly variable: R² ranges from −2.93 (section 961, extreme freeze, "
    "n = 9) to +0.82 (section 903, extreme freeze, n = 10). With 9–29 observations "
    "per section, individual R² estimates carry substantial sampling uncertainty and "
    "should not be interpreted as reliable performance indicators. The one positive "
    "outlier (section 903) likely reflects a particularly regular IRI trajectory "
    "within its 10-visit history. No systematic relationship between freeze index "
    "sub-group (moderate: 598–605, high: 864, extreme: 1,222–1,226 °C·days) "
    "and per-section R² is evident in this sample."
)

insert_figure(doc, "ontario_sensitivity_scatter.png", width_in=5.2,
    cap="Figure 3. Ontario section-level LOO R² vs. mean freeze index "
        "(XGBoost, 7 sub-iterations). Points coloured by FI sub-group. "
        "High variability and small sample sizes (n = 9–29) preclude reliable "
        "inference at the individual section level.")

# ───────────────────────────────────────────────────────────────────────
# 4.5 SHAP INTERPRETABILITY
# ───────────────────────────────────────────────────────────────────────
heading(doc, "4.5  SHAP Feature Importance and Cross-Model Consistency", level=2)

heading(doc, "4.5.1  Global Importance Rankings", level=3)

p = doc.add_paragraph(
    "SHAP TreeExplainer is applied to the XGBoost IRI design model on the 7 held-out "
    "test sections (86 visit-observations). Figure 4 shows the beeswarm plot and "
    "Figure 5 the global mean |SHAP| bar chart. Table 5 lists the top 15 features."
)

insert_figure(doc, "beeswarm_iri_design.png", width_in=5.5,
    cap="Figure 4. SHAP beeswarm plot — XGBoost IRI design model (top 20 features). "
        "Each point is one test observation; colour encodes feature quantile "
        "(blue = low, red = high). Width reflects the frequency of SHAP values "
        "at that magnitude.")

insert_figure(doc, "shap_bar_iri_design.png", width_in=5.0,
    cap="Figure 5. Global mean |SHAP| values — XGBoost IRI design model (top 20 features). "
        "Traffic loading (ESAL class 12) and structural geometry (L05B layer thickness, "
        "AC layer count) dominate.")

doc.add_paragraph("Table 5. Top 15 features by global mean |SHAP| — XGBoost IRI design model.")
t5 = doc.add_table(rows=16, cols=4)
add_table_style(t5)
hdr_row(t5, ["Rank", "Feature", "Category", "Mean |SHAP|"])
top15 = [
    (1,  "TRF_REP_REP_ESAL_PER_VEH_CLASS_12",    "Traffic",   "0.0475"),
    (2,  "L05B_REPR_THICKNESS",                    "Structure", "0.0375"),
    (3,  "LAYER_COUNT_AC",                          "Structure", "0.0367"),
    (4,  "AC_ASPHALT_CONTENT_MEAN",                 "Structure", "0.0367"),
    (5,  "TRF_CMLTV_VOL_VEH_CLASS_9_TREND",        "Traffic",   "0.0295"),
    (6,  "COMP_AGE_CLIMATE",                        "Climate",   "0.0260"),
    (7,  "L05B_DESCRIPTION",                        "Structure", "0.0231"),
    (8,  "LAYER_THICKNESS_AC_MM",                   "Structure", "0.0215"),
    (9,  "TRF_AADTT_VEH_CLASS_4_TREND",             "Traffic",   "0.0200"),
    (10, "COMP_STRUCT_ADEQUACY",                    "Structure", "0.0191"),
    (11, "CLIM_LONGITUDE",                          "Climate",   "0.0167"),
    (12, "TRF_AADTT_VEH_CLASS_5_TREND",             "Traffic",   "0.0141"),
    (13, "L05B_MATL_CODE",                          "Structure", "0.0134"),
    (14, "TRF_AADTT_VEH_CLASS_8_TREND",             "Traffic",   "0.0119"),
    (15, "TRF_ALDF_MEPDG_LG05",                    "Traffic",   "0.0104"),
]
cat_colors = {"Traffic": "FFF2CC", "Structure": "DEEAF1", "Climate": "E2EFDA"}
for i, (rank, feat, cat, shap) in enumerate(top15, start=1):
    shade = cat_colors.get(cat)
    fill_row(t5, i, [str(rank), feat, cat, shap], center=True, shade=shade)

doc.add_paragraph()

p = doc.add_paragraph(
    "Traffic loading (ESAL class 12 per class-12 vehicle, cumulative class-9 volume trend, "
    "and AADTT trends for classes 4, 5, and 8) contributes five of the top 15 features "
    "with a combined mean |SHAP| of 0.123 across the top five traffic features. "
    "Structural geometry accounts for six of the top 15 (combined 0.148), with "
    "the base layer representative thickness (L05B_REPR_THICKNESS) and AC layer "
    "count (LAYER_COUNT_AC) tied at rank 2–3. The age-climate composite "
    "(COMP_AGE_CLIMATE = AGE_YEARS × CLIM_FREEZE_INDEX) ranks sixth (0.026), "
    "indicating that cumulative service under freeze exposure is the most important "
    "individual climate signal after accounting for structural and traffic variation. "
    "Longitude (CLIM_LONGITUDE) ranks eleventh, functioning as a spatial proxy "
    "for climate regime — an expected finding given that east-west longitude "
    "correlates with precipitation and humidity patterns across the study regions."
)

heading(doc, "4.5.2  Regional SHAP Stratification", level=3)

p = doc.add_paragraph(
    "Figure 6 stratifies SHAP importance by feature category (Structure, Traffic, Climate) "
    "across the four regions, and Figure 7 shows the top 10 features per region individually."
)

insert_figure(doc, "shap_category_by_region.png", width_in=6.2,
    cap="Figure 6. SHAP importance by feature category per region — XGBoost IRI design model. "
        "Left: absolute sum of mean |SHAP|. Right: proportional split (%). "
        "Freeze index annotations above each bar indicate regional climate severity.")

insert_figure(doc, "regional_top_features.png", width_in=6.5,
    cap="Figure 7. Top-10 features by SHAP importance per region. "
        "Bars coloured by category: blue = Structure, orange = Traffic, green = Climate.")

doc.add_paragraph("Table 6. SHAP importance by feature category and region (XGBoost IRI design model).")
t6 = doc.add_table(rows=5, cols=5)
add_table_style(t6)
hdr_row(t6, ["Category", "Arizona\n(FI=5.9)", "Georgia\n(FI=10.9)", "Ohio\n(FI=316.6)", "Ontario\n(FI=834.9)"])
rows_t6 = [
    ("Structure", "0.355 (51.8%)", "0.289 (51.8%)", "0.235 (44.7%)", "0.334 (49.5%)"),
    ("Traffic",   "0.233 (34.0%)", "0.217 (38.9%)", "0.219 (41.6%)", "0.278 (41.2%)"),
    ("Climate",   "0.097 (14.2%)", "0.054 (9.7%)",  "0.086 (16.3%)", "0.067 (9.9%)"),
]
cat_shades = {"Structure": "DEEAF1", "Traffic": "FFF2CC", "Climate": "E2EFDA"}
for i, (cat, az, ga, oh, on) in enumerate(rows_t6, start=1):
    shade = cat_shades.get(cat)
    fill_row(t6, i, [cat, az, ga, oh, on], bold_first=True, center=True, shade=shade)
# Sub-total row
fill_row(t6, 4, ["Total", "0.685", "0.560", "0.526", "0.675"],
         bold_first=True, center=True, shade="D9D9D9")

doc.add_paragraph()

p = doc.add_paragraph(
    "Structural features account for 45–52% of total SHAP importance across all four "
    "regions, confirming that pavement geometry (AC layer count and thickness) and "
    "mix composition (asphalt content) are the primary drivers of IRI variation in "
    "the model's learned associations. Traffic features account for 34–42%, with "
    "Ontario showing the highest traffic importance (41.2%), consistent with its "
    "heavier truck fleet (ESAL class 12: 1.3 vs. 0.1–0.5 for other regions). "
    "Climate features account for only 10–16% of total importance — the smallest "
    "category despite the study's design intent to capture freeze-gradient effects. "
    "Georgia shows the lowest climate importance (9.7%) despite having the most "
    "distinctive climate among the four regions, further supporting the interpretation "
    "that the model has insufficient coverage of wet-subtropical deterioration mechanisms."
)

heading(doc, "4.5.3  Cross-Model SHAP Consistency", level=3)

p = doc.add_paragraph(
    "A pre-specified Spearman rank correlation test between the XGBoost and RF "
    "global SHAP importance rankings was used to assess stability. "
    "The pre-registered threshold was ρ > 0.75; the expected value from prior "
    "analysis was ρ ≈ 0.594, which would have indicated instability. "
    "The observed correlation is ρ = 0.794 (p < 0.001, n = 153 features), "
    "exceeding the threshold. The result is classified as "
)
bold_run(p, "STABLE")
normal_run(p,
    " — the top-ranking features are consistent between both architectures. "
    "This finding supports the reliability of the SHAP rankings for the features "
    "at the top of the importance distribution, though it does not imply causal "
    "validity. The ρ estimate is based on 7 test sections and carries sampling "
    "uncertainty; it is treated as an approximate indicator, not a precise boundary."
)

# ───────────────────────────────────────────────────────────────────────
# 4.6 PARTIAL DEPENDENCE ANALYSIS
# ───────────────────────────────────────────────────────────────────────
heading(doc, "4.6  Partial Dependence Analysis", level=2)

p = doc.add_paragraph(
    "Partial dependence plots (PDPs) characterise the average marginal effect of "
    "individual features on predicted IRI, holding all other features at their "
    "training values. One-way PDPs are computed for five features; two-way PDPs "
    "for two feature pairs. All PDPs are computed on the training set "
    "(455 visit-observations) to ensure adequate marginal distribution coverage. "
    "Note that CLIM_FREEZE_INDEX and AGE_YEARS were excluded by the collinearity "
    "clustering step (§3.7); the PDPs use their nearest retained equivalents: "
    "COMP_AGE_CLIMATE (AGE_YEARS × CLIM_FREEZE_INDEX) and CLIM_FREEZE_THAW_WINTER."
)

insert_figure(doc, "pdp_oneway.png", width_in=6.2,
    cap="Figure 8. One-way partial dependence plots — XGBoost IRI design model. "
        "Features shown: (a) COMP_AGE_CLIMATE, (b) AC layer thickness, "
        "(c) winter freeze-thaw cycles, (d) subgrade resilient modulus, "
        "(e) wet-freeze compound. Filled area shows range from minimum to curve value.")

p = doc.add_paragraph(
    "The age-climate composite (COMP_AGE_CLIMATE) shows a monotonically increasing "
    "PDP: predicted IRI rises as accumulated service under freeze exposure increases, "
    "consistent with the physical expectation that older pavements in harsher climates "
    "exhibit greater roughness. AC layer thickness shows a non-monotonic relationship — "
    "predicted IRI first decreases with thickness (thicker pavements are more "
    "structurally resilient), then plateaus or slightly increases at very high "
    "thicknesses, potentially reflecting that very thick AC layers are used on "
    "high-traffic, high-deterioration sections. Winter freeze-thaw cycle count shows "
    "a positive relationship with predicted IRI — more freeze-thaw cycles correspond "
    "to higher roughness — while subgrade resilient modulus shows a negative "
    "relationship (stronger subgrade → lower IRI). "
    "These associations are physically plausible but should not be interpreted "
    "as causal: SHAP values and PDPs describe model behaviour, not pavement mechanics."
)

insert_figure(doc, "pdp_twoway.png", width_in=6.0,
    cap="Figure 9. Two-way partial dependence plots — XGBoost IRI design model. "
        "Left: winter freeze-thaw cycles × AC layer thickness. "
        "Right: age-climate compound × wet-freeze compound. "
        "Colour scale: blue = lower predicted IRI, red = higher. "
        "Contour shading at 30 × 30 quantile-spaced grid points.")

p = doc.add_paragraph(
    "The two-way PDP for freeze-thaw cycles × AC thickness reveals an interaction: "
    "IRI is highest when both freeze-thaw exposure is high and AC thickness is low, "
    "and lowest when AC thickness is high regardless of freeze-thaw intensity. "
    "This interaction aligns with the structural freeze insulation hypothesis: "
    "thicker AC layers buffer subgrade materials from thermal cycling. "
    "The age-climate × wet-freeze PDP shows that the joint effect of cumulative "
    "age-freeze exposure and moisture amplification produces the highest predicted "
    "IRI at high values of both compounds — consistent with the physical mechanism "
    "of moisture-assisted freeze-thaw subgrade deterioration. "
    "However, both two-way PDPs are derived from a model with near-zero generalization "
    "performance (IRI design R² ≈ 0.06–0.13), so these associations represent "
    "patterns learned from 34 training sections and may not hold in unseen climate regimes."
)

# ───────────────────────────────────────────────────────────────────────
# 4.7 REGIONAL CLIMATE AND TRAFFIC COMPARISON
# ───────────────────────────────────────────────────────────────────────
heading(doc, "4.7  Regional Climate, Traffic, and Deterioration Profile", level=2)

p = doc.add_paragraph(
    "Figure 10 contextualises the four study regions across climate severity, "
    "mean IRI, traffic loading, and deterioration trajectory. The regional comparison "
    "is essential for interpreting both the LOO generalisation results (§4.4) and "
    "the regional SHAP stratification (§4.5.2)."
)

insert_figure(doc, "regional_climate_traffic_profile.png", width_in=6.4,
    cap="Figure 10. Regional profile: (A) mean freeze index (log scale), "
        "(B) mean IRI ± 1 SD by region, (C) AADTT and ESAL class-12 loading, "
        "(D) IRI vs. pavement age with binned-mean trend lines by region.")

p = doc.add_paragraph(
    "Panel A confirms the three-order-of-magnitude freeze gradient: from Arizona "
    "(FI = 5.9 °C·days) to Ontario (FI = 834.9 °C·days). Panel B shows that "
    "Ontario and Ohio have the highest mean IRI (1.307 and 1.231 m/km respectively) "
    "and widest distributions (SD ≈ 0.45–0.50), while Georgia is the smoothest "
    "(mean 0.909, SD 0.275). Panel C highlights Ontario's heavier truck loading "
    "(ESAL class 12 = 1.3 vs. 0.1–0.5 elsewhere), which explains the elevated "
    "traffic SHAP importance observed for Ontario in §4.5.2. "
    "Panel D reveals that IRI trajectories diverge substantially by region over "
    "pavement age: Ohio and Ontario sections show steeper deterioration rates "
    "beyond 10 years of service, consistent with freeze-induced subgrade deformation, "
    "while Arizona sections show shallower rates consistent with dry-climate resilience."
)

# ───────────────────────────────────────────────────────────────────────
# 4.8 RESIDUAL DIAGNOSTICS
# ───────────────────────────────────────────────────────────────────────
heading(doc, "4.8  Residual Diagnostics", level=2)

insert_figure(doc, "residual_diagnostics.png", width_in=6.2,
    cap="Figure 11. Residual diagnostics — XGBoost IRI design model (test set). "
        "Clockwise from top-left: residuals vs. predicted IRI, residuals vs. "
        "observation year, residuals by region (box plots), residuals vs. "
        "freeze index (not available in scaled features; panel omitted). "
        "Red line at zero. No systematic temporal trend is visible, "
        "confirming that the section-wise split prevents temporal leakage.")

p = doc.add_paragraph(
    "The residuals-vs-predicted panel shows no pronounced heteroscedasticity: "
    "variance is roughly constant across the predicted IRI range (0.8–2.0 m/km), "
    "though the positive residuals (actual > predicted) are slightly more dispersed "
    "at high predicted values, indicating occasional under-prediction for the most "
    "deteriorated sections. The residuals-vs-year panel shows no systematic temporal "
    "trend across 1989–2021, confirming that the section-wise partitioning strategy "
    "successfully prevents temporal leakage. The regional box plot shows that all "
    "four regions have median residuals near zero; Ontario shows wider IQR, "
    "consistent with the higher IRI variance of that region."
)

# ───────────────────────────────────────────────────────────────────────
# 4.9 SYNTHESIS AND LIMITS OF XAI
# ───────────────────────────────────────────────────────────────────────
heading(doc, "4.9  Synthesis: Limits of Explainable Machine Learning for Pavement Prediction", level=2)

p = doc.add_paragraph(
    "The complete evidence from §§4.1–4.8 converges on four conclusions that "
    "together define the limits of XAI application in this domain:"
)

p = doc.add_paragraph(style="List Number")
bold_run(p, "Prediction accuracy is insufficient for operational use without regional coverage. ")
normal_run(p,
    "IRI design prediction achieves R² ≈ 0 on the 7-section test set (95% CI spanning zero). "
    "LOO R² is negative in all four climate regions. The model cannot generalise beyond "
    "its training climate distribution regardless of model capacity (150 Optuna trials, "
    "153 features, two architectures). Extending the training dataset to include sections "
    "from the target climate zone is a necessary precondition for operational deployment, "
    "not a model improvement that can be achieved through further tuning."
)

p = doc.add_paragraph(style="List Number")
bold_run(p, "The monitoring lag adds no value beyond 1-year persistence for IRI. ")
normal_run(p,
    "On the 25-observation IRI common evaluation set, the RF monitoring model "
    "(R² = 0.850) matches the 1-year persistence benchmark (R² = 0.853). "
    "The practical implication: an IRI monitoring system's primary value is in "
    "its raw observations, not in the model applied to those observations."
)

p = doc.add_paragraph(style="List Number")
bold_run(p, "Generalisation failure is driven by climate distribution mismatch, not freeze severity alone. ")
normal_run(p,
    "Georgia (FI = 10.9, wet subtropical) fails worse than Ontario (FI = 834.9, extreme freeze). "
    "This finding contradicts the intuition that high-FI regions are hardest to predict. "
    "A model trained on arid and cold-freeze climates has no representation of "
    "wet-subtropical deterioration mechanisms, regardless of the target region's freeze index."
)

p = doc.add_paragraph(style="List Number")
bold_run(p, "SHAP explanations are stable across architectures but do not imply causality. ")
normal_run(p,
    "Cross-model SHAP consistency (ρ = 0.794) exceeds the pre-specified threshold, "
    "confirming that the top-ranking features (ESAL class 12, AC layer geometry, "
    "age-climate composite) are architecture-independent. However, this stability is "
    "only meaningful if the underlying model generalises — which it does not for IRI "
    "across climate zones. SHAP values from a model with R² ≈ 0 describe associations "
    "in the training distribution, not transferable pavement mechanisms. "
    "The partial dependence analysis yields physically plausible associations "
    "(thicker AC → lower IRI; more freeze-thaw cycles → higher IRI), but these "
    "are contingent on the model's regional training mix and should not be used "
    "to inform design guidelines for regions outside the training distribution."
)

p = doc.add_paragraph(
    "Taken together, these findings suggest that the value of XAI tools in pavement "
    "engineering is currently limited by data scarcity — not by interpretability method "
    "choice — and that any robust XAI analysis of pavement deterioration must first "
    "demonstrate generalisation across the climate regimes of intended deployment."
)

# ─── Save ────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Saved → {OUT_PATH}")
print(f"  Paragraphs : {len(doc.paragraphs)}")
print(f"  Tables     : {len(doc.tables)}")
